"""
Templates API Endpoints

REST endpoints for resume and requirements management.
"""

import re
import json
import hashlib
from pathlib import Path
from typing import Optional, Dict, Any, List
from datetime import datetime

from fastapi import APIRouter, HTTPException, UploadFile, File
from pydantic import BaseModel
import yaml
from dotenv import load_dotenv

from src.utils.profile_manager import ProfilePaths
from src.core.database import get_database

router = APIRouter()


def get_current_profile_paths() -> ProfilePaths:
    """Get ProfilePaths for the current active profile, reloading env first."""
    # Reload .env to get the latest ACTIVE_PROFILE
    load_dotenv(override=True)
    return ProfilePaths()


class ResumeContent(BaseModel):
    """Resume content response"""
    content: str
    last_modified: Optional[str] = None


class ResumeUpdate(BaseModel):
    """Resume update request"""
    content: str


class RequirementsContent(BaseModel):
    """Requirements content response"""
    content: str
    data: Optional[dict] = None
    last_modified: Optional[str] = None


class RequirementsUpdate(BaseModel):
    """Requirements update request"""
    content: str


class ValidationResult(BaseModel):
    """Validation result for a single template"""
    valid: bool
    exists: bool
    size: Optional[int] = None
    errors: Optional[list] = None


class TemplateValidation(BaseModel):
    """Template validation response"""
    resume: ValidationResult
    requirements: ValidationResult


class ATSCategoryResult(BaseModel):
    """ATS score for a single category"""
    score: int
    issues: List[str]
    recommendations: List[str]


class ATSScoreResponse(BaseModel):
    """ATS scoring response"""
    overall_score: int
    categories: Dict[str, ATSCategoryResult]
    summary: str
    top_recommendations: List[str]


class ResumeUploadResponse(BaseModel):
    """Response after uploading a resume"""
    success: bool
    content: str
    message: str


# Resume parsing models (matching Pydantic models in resume_parser.py)
class ContactInfoResponse(BaseModel):
    """Contact information from parsed resume"""
    name: str = ""
    email: str = ""
    phone: str = ""
    location: str = ""
    linkedin: str = ""
    website: str = ""


class ExperienceEntryResponse(BaseModel):
    """Work experience entry from parsed resume"""
    title: str = ""
    company: str = ""
    start_date: str = ""
    end_date: str = ""
    location: str = ""
    bullets: List[str] = []


class EducationEntryResponse(BaseModel):
    """Education entry from parsed resume"""
    degree: str = ""
    school: str = ""
    year: str = ""
    gpa: str = ""
    honors: str = ""


class ParsedResumeResponse(BaseModel):
    """Complete parsed resume structure"""
    contact: ContactInfoResponse
    summary: str = ""
    experience: List[ExperienceEntryResponse] = []
    education: List[EducationEntryResponse] = []
    skills: List[str] = []
    certifications: List[str] = []
    languages: List[str] = []


def get_file_info(path: Path) -> tuple[bool, Optional[int], Optional[str]]:
    """Get file existence, size, and last modified time."""
    if path.exists():
        stat = path.stat()
        return True, stat.st_size, str(stat.st_mtime)
    return False, None, None


def compute_resume_hash(content: str) -> str:
    """Compute a hash of the resume content for caching."""
    # Normalize whitespace to avoid cache misses due to minor formatting changes
    normalized = ' '.join(content.split())
    return hashlib.sha256(normalized.encode('utf-8')).hexdigest()[:32]


@router.get("/resume", response_model=ResumeContent)
async def get_resume():
    """Get current resume content."""
    paths = get_current_profile_paths()
    resume_path = paths.resume_path

    if not resume_path.exists():
        raise HTTPException(status_code=404, detail="Resume file not found")

    content = resume_path.read_text(encoding='utf-8')
    exists, size, mtime = get_file_info(resume_path)

    return ResumeContent(
        content=content,
        last_modified=mtime,
    )


@router.put("/resume")
async def update_resume(update: ResumeUpdate):
    """Update resume content."""
    paths = get_current_profile_paths()
    resume_path = paths.resume_path

    # Ensure templates directory exists
    resume_path.parent.mkdir(parents=True, exist_ok=True)

    # Write content
    resume_path.write_text(update.content, encoding='utf-8')

    return {"success": True, "message": "Resume updated"}


@router.get("/requirements", response_model=RequirementsContent)
async def get_requirements():
    """Get current requirements content."""
    paths = get_current_profile_paths()
    req_path = paths.requirements_path

    if not req_path.exists():
        raise HTTPException(status_code=404, detail="Requirements file not found")

    content = req_path.read_text(encoding='utf-8')
    exists, size, mtime = get_file_info(req_path)

    # Try to parse YAML
    data = None
    try:
        data = yaml.safe_load(content)
    except yaml.YAMLError:
        pass

    return RequirementsContent(
        content=content,
        data=data,
        last_modified=mtime,
    )


@router.put("/requirements")
async def update_requirements(update: RequirementsUpdate):
    """Update requirements content."""
    paths = get_current_profile_paths()
    req_path = paths.requirements_path

    # Validate YAML syntax
    try:
        yaml.safe_load(update.content)
    except yaml.YAMLError as e:
        raise HTTPException(
            status_code=400,
            detail=f"Invalid YAML syntax: {str(e)}"
        )

    # Ensure templates directory exists
    req_path.parent.mkdir(parents=True, exist_ok=True)

    # Write content
    req_path.write_text(update.content, encoding='utf-8')

    return {"success": True, "message": "Requirements updated"}


@router.post("/validate", response_model=TemplateValidation)
async def validate_templates():
    """Validate both templates."""
    paths = get_current_profile_paths()

    # Validate resume
    resume_path = paths.resume_path
    resume_exists, resume_size, _ = get_file_info(resume_path)
    resume_errors = []

    if not resume_exists:
        resume_errors.append("Resume file does not exist")
    elif resume_size == 0:
        resume_errors.append("Resume file is empty")
    elif resume_size < 100:
        resume_errors.append("Resume appears too short (< 100 bytes)")

    resume_valid = len(resume_errors) == 0

    # Validate requirements
    req_path = paths.requirements_path
    req_exists, req_size, _ = get_file_info(req_path)
    req_errors = []

    if not req_exists:
        req_errors.append("Requirements file does not exist")
    elif req_size == 0:
        req_errors.append("Requirements file is empty")
    else:
        try:
            content = req_path.read_text(encoding='utf-8')
            data = yaml.safe_load(content)

            if not isinstance(data, dict):
                req_errors.append("Requirements must be a YAML dictionary")
            else:
                # Check for recommended sections
                if 'candidate_profile' not in data:
                    req_errors.append("Missing 'candidate_profile' section (recommended)")
                if 'job_requirements' not in data:
                    req_errors.append("Missing 'job_requirements' section (recommended)")

        except yaml.YAMLError as e:
            req_errors.append(f"Invalid YAML syntax: {str(e)}")

    req_valid = len([e for e in req_errors if "recommended" not in e.lower()]) == 0

    return TemplateValidation(
        resume=ValidationResult(
            valid=resume_valid,
            exists=resume_exists,
            size=resume_size,
            errors=resume_errors if resume_errors else None,
        ),
        requirements=ValidationResult(
            valid=req_valid,
            exists=req_exists,
            size=req_size,
            errors=req_errors if req_errors else None,
        ),
    )


@router.get("/resume/ats-score", response_model=Optional[ATSScoreResponse])
async def get_cached_ats_score():
    """
    Get cached ATS score for the current resume if it exists.

    Returns the cached ATS score if the resume content hasn't changed
    since the last scoring, or null if no cached score exists.
    """
    paths = get_current_profile_paths()
    resume_path = paths.resume_path

    if not resume_path.exists():
        return None

    content = resume_path.read_text(encoding='utf-8')
    if not content.strip():
        return None

    # Compute hash of current resume content
    resume_hash = compute_resume_hash(content)

    # Look up cached score
    db = get_database()
    result = db.fetchone(
        """SELECT overall_score, categories_json, summary, top_recommendations_json
           FROM ats_scores
           WHERE resume_hash = ?
           ORDER BY created_at DESC
           LIMIT 1""",
        (resume_hash,)
    )

    if not result:
        return None

    overall_score, categories_json, summary, top_recommendations_json = result

    # Parse JSON fields
    try:
        categories_data = json.loads(categories_json)
        top_recommendations = json.loads(top_recommendations_json)

        # Convert to response model
        categories = {}
        for name, cat_data in categories_data.items():
            categories[name] = ATSCategoryResult(
                score=cat_data.get("score", 0),
                issues=cat_data.get("issues", []),
                recommendations=cat_data.get("recommendations", [])
            )

        return ATSScoreResponse(
            overall_score=overall_score,
            categories=categories,
            summary=summary,
            top_recommendations=top_recommendations
        )
    except (json.JSONDecodeError, KeyError) as e:
        # Invalid cached data, return None to trigger re-scoring
        return None


def extract_text_from_docx(file_content: bytes) -> str:
    """
    Extract plain text from a .docx file.

    Args:
        file_content: Raw bytes of the .docx file

    Returns:
        Extracted and cleaned text content
    """
    try:
        from docx import Document
        from io import BytesIO

        doc = Document(BytesIO(file_content))

        # Extract text from paragraphs
        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)

        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_text.append(cell_text)
                if row_text:
                    paragraphs.append(" | ".join(row_text))

        # Join with double newlines for readability
        content = "\n\n".join(paragraphs)

        # Clean up excessive whitespace
        content = re.sub(r'\n{3,}', '\n\n', content)
        content = re.sub(r' {2,}', ' ', content)

        return content.strip()

    except ImportError:
        raise HTTPException(
            status_code=500,
            detail="python-docx library not installed. Run: pip install python-docx"
        )
    except Exception as e:
        raise HTTPException(
            status_code=400,
            detail=f"Failed to parse .docx file: {str(e)}"
        )


@router.post("/resume/upload", response_model=ResumeUploadResponse)
async def upload_resume(file: UploadFile = File(...)):
    """
    Upload a .docx resume file, extract text, and save as resume.txt

    The uploaded file must be a .docx file. The extracted text will
    replace the current resume content.
    """
    # Validate file extension
    if not file.filename:
        raise HTTPException(status_code=400, detail="No filename provided")

    filename_lower = file.filename.lower()
    if not filename_lower.endswith('.docx'):
        raise HTTPException(
            status_code=400,
            detail="Invalid file type. Only .docx files are supported."
        )

    # Validate file size (max 10MB)
    MAX_SIZE = 10 * 1024 * 1024
    content = await file.read()
    if len(content) > MAX_SIZE:
        raise HTTPException(
            status_code=400,
            detail=f"File too large. Maximum size is 10MB."
        )

    # Extract text from docx
    extracted_text = extract_text_from_docx(content)

    if not extracted_text:
        raise HTTPException(
            status_code=400,
            detail="No text content found in the document."
        )

    # Save to resume.txt
    paths = get_current_profile_paths()
    resume_path = paths.resume_path

    # Ensure templates directory exists
    resume_path.parent.mkdir(parents=True, exist_ok=True)

    # Write content
    resume_path.write_text(extracted_text, encoding='utf-8')

    return ResumeUploadResponse(
        success=True,
        content=extracted_text,
        message=f"Resume uploaded and extracted from {file.filename}"
    )


@router.post("/resume/ats-score", response_model=ATSScoreResponse)
async def score_resume_ats():
    """
    Run ATS quality scoring on the current resume.

    Analyzes the resume for ATS compatibility including:
    - Keyword optimization
    - Formatting compatibility
    - Section structure
    - Quantified achievements
    - Contact information
    - Skills presentation

    Requires llama-server to be running.
    """
    paths = get_current_profile_paths()
    resume_path = paths.resume_path

    if not resume_path.exists():
        raise HTTPException(status_code=404, detail="Resume file not found")

    content = resume_path.read_text(encoding='utf-8')

    if not content.strip():
        raise HTTPException(status_code=400, detail="Resume is empty")

    # Import and run ATS scorer
    try:
        from src.job_matcher.ats_scorer import ATSScorer

        scorer = ATSScorer()

        # Check connection first
        if not scorer.test_connection():
            raise HTTPException(
                status_code=503,
                detail="AI server (llama-server) is not available. Please ensure it is running."
            )

        result = scorer.score_resume(content)

        if not result:
            raise HTTPException(
                status_code=500,
                detail="ATS scoring failed. Please try again."
            )

        # Convert to response model
        categories = {}
        categories_dict = {}
        for name, cat in result.categories.items():
            categories[name] = ATSCategoryResult(
                score=cat.score,
                issues=cat.issues,
                recommendations=cat.recommendations
            )
            categories_dict[name] = {
                "score": cat.score,
                "issues": cat.issues,
                "recommendations": cat.recommendations
            }

        # Save to database for caching
        resume_hash = compute_resume_hash(content)
        db = get_database()
        now = datetime.now()

        # Delete any existing score for this resume hash and insert new one
        db.execute("DELETE FROM ats_scores WHERE resume_hash = ?", (resume_hash,))
        db.execute(
            """INSERT INTO ats_scores (
                resume_hash, overall_score, categories_json, summary,
                top_recommendations_json, created_at, updated_at
            ) VALUES (?, ?, ?, ?, ?, ?, ?)""",
            (
                resume_hash,
                result.overall_score,
                json.dumps(categories_dict),
                result.summary,
                json.dumps(result.top_recommendations),
                now,
                now
            )
        )

        return ATSScoreResponse(
            overall_score=result.overall_score,
            categories=categories,
            summary=result.summary,
            top_recommendations=result.top_recommendations
        )

    except ImportError as e:
        raise HTTPException(
            status_code=500,
            detail=f"Failed to import ATS scorer: {str(e)}"
        )
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"ATS scoring error: {str(e)}"
        )


@router.post("/resume/parse", response_model=ParsedResumeResponse)
async def parse_resume():
    """
    Parse the current resume into structured data using AI.

    Uses LLM with Pydantic schema enforcement to extract:
    - Contact information (name, email, phone, location, LinkedIn)
    - Professional summary
    - Work experience with accomplishments
    - Education history
    - Skills, certifications, languages

    Requires llama-server to be running.
    """
    paths = get_current_profile_paths()
    resume_path = paths.resume_path

    if not resume_path.exists():
        raise HTTPException(status_code=404, detail="Resume file not found")

    content = resume_path.read_text(encoding='utf-8')

    if not content.strip():
        raise HTTPException(status_code=400, detail="Resume is empty")

    try:
        from src.job_matcher.resume_parser import ResumeParser

        parser = ResumeParser()

        # Check connection first
        if not parser.test_connection():
            raise HTTPException(
                status_code=503,
                detail="AI server (llama-server) is not available. Please ensure it is running."
            )

        result = parser.parse(content)

        if not result:
            raise HTTPException(
                status_code=500,
                detail="Resume parsing failed. Please try again."
            )

        # Convert to response model
        return ParsedResumeResponse(
            contact=ContactInfoResponse(
                name=result.contact.name,
                email=result.contact.email,
                phone=result.contact.phone,
                location=result.contact.location,
                linkedin=result.contact.linkedin,
                website=result.contact.website
            ),
            summary=result.summary,
            experience=[
                ExperienceEntryResponse(
                    title=exp.title,
                    company=exp.company,
                    start_date=exp.start_date,
                    end_date=exp.end_date,
                    location=exp.location,
                    bullets=exp.bullets
                )
                for exp in result.experience
            ],
            education=[
                EducationEntryResponse(
                    degree=edu.degree,
                    school=edu.school,
                    year=edu.year,
                    gpa=edu.gpa,
                    honors=edu.honors
                )
                for edu in result.education
            ],
            skills=result.skills,
            certifications=result.certifications,
            languages=result.languages
        )

    except ImportError as e:
        raise HTTPException(
            status_code=500,
            detail=f"Failed to import resume parser: {str(e)}"
        )
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Resume parsing error: {str(e)}"
        )
